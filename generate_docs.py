"""
Generate PDF and Word documents from Markdown files
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
import markdown

def markdown_to_word(md_file, output_file):
    """Convert Markdown to Word document"""
    print(f"Converting {md_file} to Word... (Step 1/2)")
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    doc.add_heading('HR Assistant System - Project Summary', 0)
    
    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), 1)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), 2)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), 3)
        elif line.startswith('#### '):
            doc.add_heading(line.replace('#### ', ''), 4)
        elif line.startswith('- '):
            doc.add_paragraph(line.replace('- ', ''), style='List Bullet')
        elif line.startswith('| '):
            # Simple table handling
            continue
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()
    
    doc.save(output_file)
    print(f"Done: Word document saved: {output_file}")

def markdown_to_pdf(md_file, output_file):
    """Convert Markdown to PDF document"""
    print(f"Converting {md_file} to PDF... (Step 2/2)")
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Convert markdown to HTML
    html = markdown.markdown(content)
    
    doc = SimpleDocTemplate(output_file, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    story = []
    
    # Add title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=RGBColor(0, 51, 102),
        spaceAfter=30,
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    story.append(Paragraph("HR Assistant System - Project Summary", title_style))
    story.append(Spacer(1, 0.5*inch))
    
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        if line.startswith('# '):
            heading_text = line.replace('# ', '')
            heading_style = ParagraphStyle(
                'CustomHeading1',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=RGBColor(0, 51, 102),
                spaceAfter=12
            )
            story.append(Paragraph(heading_text, heading_style))
            story.append(Spacer(1, 0.2*inch))
        
        elif line.startswith('## '):
            heading_text = line.replace('## ', '')
            heading_style = ParagraphStyle(
                'CustomHeading2',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=RGBColor(0, 102, 204),
                spaceAfter=10
            )
            story.append(Paragraph(heading_text, heading_style))
            story.append(Spacer(1, 0.15*inch))
        
        elif line.startswith('### '):
            heading_text = line.replace('### ', '')
            story.append(Paragraph(f"<b>{heading_text}</b>", styles['Heading3']))
            story.append(Spacer(1, 0.1*inch))
        
        elif line.startswith('- '):
            bullet_text = line.replace('- ', '')
            story.append(Paragraph(f"• {bullet_text}", styles['Normal']))
        
        elif line.startswith('```'):
            # Code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            code_style = ParagraphStyle(
                'Code',
                parent=styles['Normal'],
                fontName='Courier',
                fontSize=9,
                textColor=RGBColor(64, 64, 64),
                backColor=RGBColor(240, 240, 240),
                leftIndent=20
            )
            for code_line in code_text.split('\n'):
                if code_line.strip():
                    story.append(Paragraph(code_line, code_style))
            story.append(Spacer(1, 0.1*inch))
        
        elif line.startswith('| '):
            # Table row
            story.append(Paragraph(line, styles['Normal']))
        
        elif line.strip():
            story.append(Paragraph(line, styles['Normal']))
        
        story.append(Spacer(1, 0.05*inch))
        i += 1
    
    doc.build(story)
    print(f"Done: PDF document saved: {output_file}")

def main():
    project_dir = Path(r"c:\Users\lethu\OneDrive\Máy tính\AI\WORKSHOP 4")
    md_file = project_dir / "PROJECT_SUMMARY.md"
    
    if not md_file.exists():
        print(f"❌ File not found: {md_file}")
        return
    
    # Generate Word
    word_output = project_dir / "PROJECT_SUMMARY.docx"
    markdown_to_word(str(md_file), str(word_output))
    
    # Generate PDF
    pdf_output = project_dir / "PROJECT_SUMMARY.pdf"
    markdown_to_pdf(str(md_file), str(pdf_output))
    
    print("\n" + "="*50)
    print("SUCCESS: Documents Generated!")
    print("="*50)
    print(f"Word: {word_output.name}")
    print(f"PDF:  {pdf_output.name}")
    print("="*50)

if __name__ == "__main__":
    main()
